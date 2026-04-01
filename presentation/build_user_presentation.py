from __future__ import annotations

import datetime as _dt
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
ASSETS_DIR = ROOT / "presentation" / "assets"
OUT_PATH = ROOT / "presentation" / "RestaurantAnalyze_User_Presentation.docx"


def _set_landscape(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)


def _set_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(18)

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(40)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(32)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(24)


def _add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph(title, style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    sp = doc.add_paragraph(subtitle)
    if sp.runs:
        sp.runs[0].font.size = Pt(20)


def _add_bullets(doc: Document, bullets: list[str]) -> None:
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")


def _add_picture(doc: Document, path: Path, width_inches: float = 9.6, caption: str | None = None) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph(caption)
        if cap.runs:
            cap.runs[0].font.size = Pt(14)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _slide(
    doc: Document,
    title: str,
    bullets: list[str] | None = None,
    images: list[tuple[Path, str | None]] | None = None,
) -> None:
    doc.add_paragraph(title, style="Heading 1")
    if bullets:
        _add_bullets(doc, bullets)
    if images:
        for p, cap in images:
            _add_picture(doc, p, caption=cap)
    doc.add_page_break()


def main() -> None:
    doc = Document()
    _set_landscape(doc)
    _set_base_styles(doc)

    today = _dt.date.today().strftime("%b %d, %Y")

    _add_title(
        doc,
        "Restaurant Recommendation System",
        f"User Guide Presentation • {today}",
    )
    _add_picture(doc, ASSETS_DIR / "03_landing.png", caption="Home page")
    doc.add_page_break()

    _slide(
        doc,
        "Overview (What It Does)",
        bullets=[
            "Helps customers quickly find the best food and drink options from the restaurant’s menu.",
            "You can search by cravings (keywords), browse the full menu, chat with the assistant, and build an order in the cart.",
        ],
        images=[(ASSETS_DIR / "03_landing.png", "Start from the Home page")],
    )

    _slide(
        doc,
        "How to Use — Recommendations",
        bullets=[
            "On the Home page, type what you want to eat and/or drink (at least one field).",
            "Click “Get Recommendations”.",
            "You’ll see recommended foods and drinks with prices (and match indicators).",
            "Use “Try Again” to quickly run a new search.",
        ],
        images=[
            (ASSETS_DIR / "05_preference_form.png", "Enter your cravings"),
            (ASSETS_DIR / "06_recommendations_top.png", "Results page"),
        ],
    )

    _slide(
        doc,
        "How to Use — Explore Recommendation Cards",
        bullets=[
            "Each recommendation appears as a card with the item name, short description, and price.",
            "Use “Add to Cart” to build your order.",
            "Use “Show similar” to explore alternatives (useful if you want something close but different).",
        ],
        images=[(ASSETS_DIR / "07_recommendation_cards.png", "Recommendation cards")],
    )

    _slide(
        doc,
        "How to Use — Chef’s Picks",
        bullets=[
            "Chef’s Picks are curated highlights shown on the Home page.",
            "Great for quick browsing when you don’t want to type a search.",
            "You can add picks directly to your cart.",
        ],
        images=[(ASSETS_DIR / "04_preference_form_and_picks.png", "Chef’s Picks section")],
    )

    _slide(
        doc,
        "How to Use — Full Menu",
        bullets=[
            "Click “Full Menu” in the navigation bar.",
            "Browse foods and drinks grouped by categories.",
            "Use the “Add” buttons to add items to your cart while browsing.",
        ],
        images=[
            (ASSETS_DIR / "08_menu_top.png", "Full Menu page"),
            (ASSETS_DIR / "09_menu_table.png", "Menu browsing + quick add"),
        ],
    )

    _slide(
        doc,
        "How to Use — Cart",
        bullets=[
            "Click the cart icon (🛒) to open the cart panel.",
            "Increase/decrease quantities, remove items, or clear the cart.",
            "When ready, you can place the order (demo flow) and get an order number.",
        ],
        images=[(ASSETS_DIR / "10_cart_panel_empty.png", "Cart panel")],
    )

    _slide(
        doc,
        "How to Use — Chatbot (Concierge)",
        bullets=[
            "Click the chat bubble to open the concierge assistant.",
            "Ask questions like: “What do you recommend with kebab?” or “Do you have something sweet?”",
            "The chatbot helps you discover menu items and make choices faster.",
        ],
        images=[(ASSETS_DIR / "11_chat_panel.png", "Chat widget")],
    )

    _slide(
        doc,
        "Typical Demo Flow (3–5 minutes)",
        bullets=[
            "Home → type a craving → show recommendations.",
            "Add 2–3 items to cart → open cart panel and adjust quantities.",
            "Open Full Menu → browse categories and add one more item.",
            "Open chatbot → ask a question about a dessert or drink suggestion.",
        ],
        images=[(ASSETS_DIR / "03_landing.png", "Suggested demo starting point")],
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Wrote: {OUT_PATH}")


if __name__ == "__main__":
    main()

