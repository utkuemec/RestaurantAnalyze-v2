from __future__ import annotations

import datetime as _dt
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
ASSETS_DIR = ROOT / "presentation" / "assets"
OUT_PATH = ROOT / "presentation" / "RestaurantAnalyze_Features_With_Images.docx"


def _set_landscape(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)


def _set_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(18)

    title = doc.styles["Title"]
    title.font.name = "Calibri"
    title.font.size = Pt(40)

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(32)

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(24)


def _add_slide_title(doc: Document, title: str, subtitle: str | None = None) -> None:
    p = doc.add_paragraph(title, style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if subtitle:
        sp = doc.add_paragraph(subtitle)
        sp.runs[0].font.size = Pt(20)
        sp.runs[0].font.color.rgb = None


def _add_bullets(doc: Document, bullets: list[str]) -> None:
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")


def _add_picture(doc: Document, path: Path, width_inches: float = 9.6, caption: str | None = None) -> None:
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph(caption)
        cap.runs[0].font.size = Pt(14)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _page_break(doc: Document) -> None:
    doc.add_page_break()


def _slide(
    doc: Document,
    title: str,
    bullets: list[str] | None = None,
    images: list[tuple[Path, str | None]] | None = None,
) -> None:
    doc.add_paragraph(title, style="Heading 1")
    if bullets:
        _add_bullets(doc, bullets)
    if images:
        for p, cap in images:
            _add_picture(doc, p, caption=cap)
    _page_break(doc)


def _api_table_slide(doc: Document) -> None:
    doc.add_paragraph("Key API Endpoints", style="Heading 1")
    doc.add_paragraph("Backend (Flask, port 5008)", style="Heading 2")

    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Method"
    hdr[1].text = "Endpoint"
    hdr[2].text = "Purpose"

    rows = [
        ("GET", "/api/menu", "Full menu (foods + drinks)"),
        ("POST", "/api/recommend", "Food+drink recommendations (strategy-based)"),
        ("POST", "/api/track", "Track interactions (view/click/like/order/rate)"),
        ("GET", "/api/chefs-picks?per=3", "Chef’s picks (per category)"),
        ("GET", "/api/popular?n=8", "Trending/popular items"),
        ("GET", "/api/similar?item=...&n=6", "Similar items using item-item similarity"),
        ("GET", "/api/personalized?user_id=...&n=8", "Personalized recommendations (CF+BPR)"),
        ("POST", "/api/retrain", "Retrain CF/BPR/pairing models from interactions"),
        ("GET", "/api/stats", "Engine statistics (items/users/interactions/etc.)"),
    ]
    for m, ep, purp in rows:
        r = table.add_row().cells
        r[0].text = m
        r[1].text = ep
        r[2].text = purp

    doc.add_paragraph()
    doc.add_paragraph("Frontend (Express, port 3001)", style="Heading 2")
    table2 = doc.add_table(rows=1, cols=3)
    hdr2 = table2.rows[0].cells
    hdr2[0].text = "Route"
    hdr2[1].text = "Type"
    hdr2[2].text = "Purpose"

    rows2 = [
        ("/", "Page", "Recommend form + Chef’s Picks"),
        ("/preferences", "Page", "Recommendation results"),
        ("/menu", "Page", "Full menu browsing"),
        ("/api/like", "API", "Like an item (proxies to backend tracking)"),
        ("/api/chat", "API", "AI concierge chat (Groq/Gemini optional)"),
        ("/api/order", "API", "Create order (in-memory demo)"),
        ("/api/orders", "API", "View received orders (in-memory demo)"),
    ]
    for route, typ, purp in rows2:
        r = table2.add_row().cells
        r[0].text = route
        r[1].text = typ
        r[2].text = purp

    _page_break(doc)


def main() -> None:
    doc = Document()
    _set_landscape(doc)
    _set_base_styles(doc)

    today = _dt.date.today().strftime("%b %d, %Y")

    _add_slide_title(
        doc,
        "Restaurant Recommendation System",
        f"Feature Overview & Demo Walkthrough • {today}",
    )
    _add_picture(doc, ASSETS_DIR / "03_landing.png", caption="Landing page (Recommend)")
    _page_break(doc)

    _slide(
        doc,
        "What the System Does",
        bullets=[
            "Takes a customer’s food/drink cravings and recommends the best matching menu items.",
            "Supports multiple recommendation strategies (content, fuzzy, CF, BPR, popularity, hybrid).",
            "Provides a modern web UI with menu browsing, cart panel, and optional AI concierge chat.",
            "Learns from interactions (likes/orders) to improve personalization over time.",
        ],
        images=[(ASSETS_DIR / "12_architecture.png", "High-level architecture")],
    )

    _slide(
        doc,
        "User Experience — Recommend Flow",
        bullets=[
            "Customer enters a food and/or drink keyword (examples: kebab, coffee, lemonade).",
            "System returns Top-N foods and drinks with match scores.",
            "Users can explore similar items and add recommendations to the cart.",
        ],
        images=[
            (ASSETS_DIR / "05_preference_form.png", "Preference input form"),
            (ASSETS_DIR / "06_recommendations_top.png", "Recommendation results (top)"),
            (ASSETS_DIR / "07_recommendation_cards.png", "Cards show price + match score + add-to-cart"),
        ],
    )

    _slide(
        doc,
        "Chef’s Picks (Home Page)",
        bullets=[
            "Shows curated picks by category (e.g., appetizers/sides, mains, desserts, drinks).",
            "Each pick includes description and price and can be added to the cart.",
        ],
        images=[(ASSETS_DIR / "04_preference_form_and_picks.png", "Chef’s Picks section")],
    )

    _slide(
        doc,
        "Full Menu Browsing",
        bullets=[
            "Browse all foods and drinks grouped by category.",
            "Quick “Add” buttons support faster ordering exploration.",
        ],
        images=[
            (ASSETS_DIR / "08_menu_top.png", "Full Menu page"),
            (ASSETS_DIR / "09_menu_table.png", "Menu items with quick add buttons"),
        ],
    )

    _slide(
        doc,
        "Cart Panel (Order Builder)",
        bullets=[
            "Sticky cart icon opens a side panel.",
            "Supports quantity changes, remove items, clear cart, and order submission (demo).",
            "Cart is stored in localStorage for a smooth UX across pages.",
        ],
        images=[(ASSETS_DIR / "10_cart_panel_empty.png", "Cart panel UI (example)")],
    )

    _slide(
        doc,
        "AI Concierge Chat (Optional)",
        bullets=[
            "Chat widget answers questions about menu items and makes recommendations.",
            "Uses only real menu items (no hallucinated items).",
            "Language-matching: replies in the same language as the customer message.",
            "Providers: Groq (Llama) if configured, otherwise Gemini (optional).",
        ],
        images=[(ASSETS_DIR / "11_chat_panel.png", "Chat widget entry point")],
    )

    _slide(
        doc,
        "Recommendation Engine — Strategies",
        bullets=[
            "Content-based: TF‑IDF over item name/category/description + cosine similarity.",
            "Fuzzy matching: Levenshtein ratio + WordNet synonym expansion.",
            "Collaborative filtering: user×item implicit matrix → Truncated SVD.",
            "BPR: pairwise ranking optimized with SGD for implicit feedback.",
            "Popularity: time-decayed interaction-weighted trending score.",
            "Hybrid: weighted ensemble that redistributes weights if models are not trained yet.",
        ],
        images=[(ASSETS_DIR / "13_algorithms.png", "Hybrid ensemble overview")],
    )

    _slide(
        doc,
        "Learning & Personalization",
        bullets=[
            "Tracks interactions such as view/click/like/order to build implicit ratings.",
            "Persists events to `user_data/interactions.json` (simple, portable storage).",
            "Cold-start seeding generates synthetic profiles so CF/BPR work immediately.",
            "Retrain endpoint updates CF/BPR + drink pairing model from recent interactions.",
        ],
    )

    _api_table_slide(doc)

    _slide(
        doc,
        "Demo Script (5–7 minutes)",
        bullets=[
            "Open Home → enter a craving (food/drink) → show recommendations + match scores.",
            "Open Full Menu → add a few items → open cart panel.",
            "Show Like/interaction behavior (improves popularity/personalization over time).",
            "Optionally show the AI concierge answering a menu question (if API key configured).",
        ],
        images=[(ASSETS_DIR / "03_landing.png", "Start here")],
    )

    # Remove last trailing page break for cleaner export
    if doc.paragraphs and doc.paragraphs[-1].text == "":
        pass

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Wrote: {OUT_PATH}")


if __name__ == "__main__":
    main()

